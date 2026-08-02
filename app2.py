from __future__ import annotations  # compatibilidad de anotaciones de tipo con Python < 3.9

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from io import BytesIO
from pathlib import Path
from datetime import datetime
from collections import defaultdict
import json
import sqlite3
import shutil

# ====================================
# DEPENDENCIAS OPCIONALES
# ====================================

try:
    from fpdf import FPDF
    FPDF_DISPONIBLE = True
except ImportError:
    FPDF_DISPONIBLE = False

try:
    import docx  # python-docx
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

try:
    import pdfplumber
    PDFPLUMBER_DISPONIBLE = True
except ImportError:
    PDFPLUMBER_DISPONIBLE = False

try:
    from openpyxl.drawing.image import Image as OpenpyxlImage
    OPENPYXL_IMAGENES_DISPONIBLE = True
except ImportError:
    OPENPYXL_IMAGENES_DISPONIBLE = False

try:
    import kaleido  # noqa: F401  (necesario para exportar gráficos de plotly a imagen)
    KALEIDO_DISPONIBLE = True
except ImportError:
    KALEIDO_DISPONIBLE = False


# ====================================
# CONFIGURACIÓN DE LA PÁGINA
# ====================================

st.set_page_config(
    page_title="Plataforma de Reportes",
    page_icon="📊",
    layout="wide"
)

st.title("📊 Plataforma de Reportes y Análisis de Datos")

st.markdown("""
Esta plataforma permite cargar **varios archivos a la vez**, en distintos formatos
(Excel, CSV, Word, PDF y bases de datos SQLite), analizarlos, generar gráficos
y exportar reportes que incluyen los gráficos generados.
""")

with st.expander("GUIA RAPIDA DE USO"):
    st.markdown("""
    1. Suba uno o **varios** archivos: **.xlsx**, **.xls**, **.csv**, **.docx**, **.pdf**, **.db/.sqlite**.
    2. Si un archivo tiene varias hojas/tablas (Excel, Word, PDF, base de datos), elija cuáles desea cargar;
       cada una se convierte en un "dataset" independiente que puede analizar o graficar.
    3. En **"Datasets a incluir en el análisis y el reporte"** elija **uno o varios datasets a la vez**
       (por ejemplo, uno de un archivo Excel y otro de un archivo CSV distinto). Cada dataset tiene su
       propia pestaña con sus propios filtros e indicadores.
    4. Revise las advertencias de calidad de datos antes de interpretar los indicadores.
    5. Agregue uno o **varios gráficos**, cada uno puede usar un dataset distinto (se grafican los datos
       ya filtrados).
    6. Descargue el reporte en Excel, CSV o PDF. El reporte **combina los datos filtrados de todos los
       datasets elegidos en el paso 3** (una hoja por dataset en Excel, una sección por dataset en PDF),
       y los archivos Excel y PDF **incluyen los gráficos** generados, con **todas las filas** (no solo una muestra).
       En el PDF, el o los gráficos de cada dataset aparecen justo después de su tabla de datos filtrados.
    7. Su sesión se guarda automáticamente. Si cierra la app, al volver a abrirla puede pulsar
       **"Cargar última sesión"** para recuperar los archivos y gráficos en los que estaba trabajando.
    8. Además, cada vez que presiona **"Guardar ahora"** se crea una **nueva entrada en el historial**
       (con fecha, hora, minuto y segundo). Puede volver a cargar cualquiera de esos momentos guardados
       desde la barra lateral, en **"Historial de guardados"**.

    **Columnas admitidas:** cualquier nombre de columna funciona; no es necesario que se llame
    exactamente "Categoria" o "Total". El sistema detecta automáticamente columnas categóricas,
    numéricas y de fecha, incluso si tienen espacios extra o mayúsculas distintas.
    """)

if not FPDF_DISPONIBLE or not DOCX_DISPONIBLE or not PDFPLUMBER_DISPONIBLE or not KALEIDO_DISPONIBLE:
    faltantes = []
    if not FPDF_DISPONIBLE:
        faltantes.append("fpdf2 (exportar PDF)")
    if not DOCX_DISPONIBLE:
        faltantes.append("python-docx (leer archivos Word)")
    if not PDFPLUMBER_DISPONIBLE:
        faltantes.append("pdfplumber (leer archivos PDF)")
    if not KALEIDO_DISPONIBLE:
        faltantes.append("kaleido (incluir gráficos en Excel/PDF)")
    st.caption(
        "⚠️ Algunas funciones están limitadas porque faltan librerías: "
        + ", ".join(faltantes)
        + ". Instálelas con pip para habilitarlas por completo."
    )


# ====================================
# PERSISTENCIA DE SESIÓN
# ====================================
#
# Hay dos mecanismos independientes:
#
# 1) AUTOGUARDADO ("última sesión"): se sobrescribe automáticamente en varios
#    puntos del flujo (al subir archivos, al agregar/quitar gráficos, al final
#    de cada ejecución) para poder recuperar el trabajo si se cierra la app.
#
# 2) HISTORIAL DE GUARDADOS MANUALES: se crea una entrada NUEVA (no se
#    sobrescribe ninguna anterior) cada vez que el usuario presiona el botón
#    "Guardar ahora". Cada entrada queda con su fecha y hora exacta
#    (día/mes/año hora:minuto:segundo) y puede cargarse de forma independiente.

CACHE_DIR = Path(".cache_reportes")
CACHE_ARCHIVOS_DIR = CACHE_DIR / "archivos"
META_PATH = CACHE_DIR / "sesion.json"

HISTORIAL_DIR = CACHE_DIR / "historial"


def hay_sesion_guardada() -> bool:
    return META_PATH.exists()


def guardar_sesion():
    """Autoguardado: sobrescribe la 'última sesión' con el estado actual."""
    try:
        CACHE_ARCHIVOS_DIR.mkdir(parents=True, exist_ok=True)
        for nombre, info in st.session_state.archivos_originales.items():
            ruta = CACHE_ARCHIVOS_DIR / nombre
            if not ruta.exists():
                ruta.write_bytes(info["bytes"])

        meta = {
            "archivos": [
                {"nombre": nombre, "tipo": info["tipo"]}
                for nombre, info in st.session_state.archivos_originales.items()
            ],
            "graficos": st.session_state.graficos,
            "datasets_seleccionados": st.session_state.get("datasets_seleccionados", []),
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        META_PATH.write_text(json.dumps(meta, ensure_ascii=False, indent=2))
    except Exception as e:
        st.sidebar.warning(f"No fue posible guardar la sesión: {e}")


def cargar_sesion():
    """Carga la 'última sesión' autoguardada."""
    if not META_PATH.exists():
        return
    try:
        meta = json.loads(META_PATH.read_text())
    except Exception:
        return

    for archivo_meta in meta.get("archivos", []):
        nombre = archivo_meta["nombre"]
        ruta = CACHE_ARCHIVOS_DIR / nombre
        if ruta.exists():
            st.session_state.archivos_originales[nombre] = {
                "bytes": ruta.read_bytes(),
                "tipo": archivo_meta["tipo"],
            }
    st.session_state.graficos = meta.get("graficos", [])
    st.session_state.datasets_seleccionados = meta.get("datasets_seleccionados", [])


def borrar_sesion_guardada():
    """Borra únicamente el autoguardado de 'última sesión' (no toca el historial)."""
    try:
        if META_PATH.exists():
            META_PATH.unlink()
        if CACHE_ARCHIVOS_DIR.exists():
            for f in CACHE_ARCHIVOS_DIR.glob("*"):
                f.unlink()
    except Exception as e:
        st.sidebar.warning(f"No fue posible borrar la sesión guardada: {e}")


# ---- Historial de guardados manuales ----

def guardar_historial() -> str | None:
    """
    Crea una NUEVA entrada en el historial de guardados (nunca sobrescribe las
    anteriores). Se usa exclusivamente cuando el usuario presiona "Guardar ahora".
    El identificador de la entrada incluye fecha, hora, minuto, segundo y
    microsegundos (para evitar colisiones si se guarda dos veces muy seguido).
    """
    try:
        HISTORIAL_DIR.mkdir(parents=True, exist_ok=True)
        ahora = datetime.now()
        id_guardado = ahora.strftime("%Y%m%d_%H%M%S_%f")
        carpeta = HISTORIAL_DIR / id_guardado
        carpeta_archivos = carpeta / "archivos"
        carpeta_archivos.mkdir(parents=True, exist_ok=True)

        for nombre, info in st.session_state.archivos_originales.items():
            (carpeta_archivos / nombre).write_bytes(info["bytes"])

        meta = {
            "id": id_guardado,
            "fecha": ahora.strftime("%d/%m/%Y %H:%M:%S"),
            "archivos": [
                {"nombre": nombre, "tipo": info["tipo"]}
                for nombre, info in st.session_state.archivos_originales.items()
            ],
            "graficos": st.session_state.graficos,
            "datasets_seleccionados": st.session_state.get("datasets_seleccionados", []),
        }
        (carpeta / "meta.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2))
        return id_guardado
    except Exception as e:
        st.sidebar.warning(f"No fue posible guardar en el historial: {e}")
        return None


def listar_historial() -> list:
    """Devuelve la lista de guardados manuales, del más reciente al más antiguo."""
    if not HISTORIAL_DIR.exists():
        return []
    entradas = []
    for carpeta in HISTORIAL_DIR.iterdir():
        meta_path = carpeta / "meta.json"
        if meta_path.exists():
            try:
                entradas.append(json.loads(meta_path.read_text()))
            except Exception:
                continue
    entradas.sort(key=lambda m: m.get("id", ""), reverse=True)
    return entradas


def cargar_desde_historial(id_guardado: str):
    """Reemplaza el estado actual por el de una entrada específica del historial."""
    carpeta = HISTORIAL_DIR / id_guardado
    meta_path = carpeta / "meta.json"
    if not meta_path.exists():
        return
    try:
        meta = json.loads(meta_path.read_text())
    except Exception:
        return

    st.session_state.archivos_originales = {}
    for archivo_meta in meta.get("archivos", []):
        nombre = archivo_meta["nombre"]
        ruta = carpeta / "archivos" / nombre
        if ruta.exists():
            st.session_state.archivos_originales[nombre] = {
                "bytes": ruta.read_bytes(),
                "tipo": archivo_meta["tipo"],
            }
    st.session_state.graficos = meta.get("graficos", [])
    st.session_state.datasets_seleccionados = meta.get("datasets_seleccionados", [])


def borrar_entrada_historial(id_guardado: str):
    try:
        carpeta = HISTORIAL_DIR / id_guardado
        if carpeta.exists():
            shutil.rmtree(carpeta)
    except Exception as e:
        st.sidebar.warning(f"No fue posible borrar ese guardado: {e}")


def borrar_historial_completo():
    try:
        if HISTORIAL_DIR.exists():
            shutil.rmtree(HISTORIAL_DIR)
    except Exception as e:
        st.sidebar.warning(f"No fue posible borrar el historial: {e}")


# ====================================
# ESTADO DE SESIÓN (multi-archivo, multi-dataset, multi-gráfico)
# ====================================

if "archivos_originales" not in st.session_state:
    st.session_state.archivos_originales = {}   # nombre_archivo -> {"bytes":..., "tipo":...}
if "datasets" not in st.session_state:
    st.session_state.datasets = {}               # nombre_dataset -> DataFrame procesado
if "graficos" not in st.session_state:
    st.session_state.graficos = []                # lista de dicts {id, dataset, eje_x, eje_y, tipo}
if "datasets_seleccionados" not in st.session_state:
    st.session_state.datasets_seleccionados = []   # lista de nombres de dataset incluidos en el reporte


# Ofrecer recuperar la sesión anterior si aún no hay archivos cargados en esta sesión
if not st.session_state.archivos_originales and hay_sesion_guardada():
    with st.sidebar:
        st.markdown("### 🕒 Sesión anterior detectada")
        try:
            info_meta = json.loads(META_PATH.read_text())
            st.caption(
                f"Guardada: {info_meta.get('fecha', '?')} — "
                f"{len(info_meta.get('archivos', []))} archivo(s)"
            )
        except Exception:
            pass
        if st.button(" Cargar última sesión"):
            cargar_sesion()
            st.rerun()

with st.sidebar:
    st.markdown("###  Gestión de sesión")
    col_guardar, col_borrar = st.columns(2)
    with col_guardar:
        if st.button("Guardar ahora"):
            guardar_sesion()          # actualiza la 'última sesión'
            id_nuevo = guardar_historial()  # crea una entrada nueva en el historial
            if id_nuevo:
                st.success("Sesión guardada y agregada al historial.")
    with col_borrar:
        if st.button("Borrar guardada"):
            borrar_sesion_guardada()
            st.success("Sesión guardada eliminada.")

    st.markdown("###  Historial de guardados")
    historial = listar_historial()
    if not historial:
        st.caption("Todavía no hay guardados en el historial. Use 'Guardar ahora' para crear el primero.")
    else:
        st.caption(f"{len(historial)} guardado(s) disponible(s).")
        for entrada in historial:
            n_archivos = len(entrada.get("archivos", []))
            with st.expander(f"🕓 {entrada.get('fecha', '?')}  ({n_archivos} archivo(s))"):
                nombres_archivos = ", ".join(a["nombre"] for a in entrada.get("archivos", []))
                if nombres_archivos:
                    st.caption(nombres_archivos)
                col_cargar, col_eliminar = st.columns(2)
                with col_cargar:
                    if st.button(" Cargar", key=f"cargar_hist_{entrada['id']}"):
                        cargar_desde_historial(entrada["id"])
                        st.rerun()
                with col_eliminar:
                    if st.button(" Eliminar", key=f"borrar_hist_{entrada['id']}"):
                        borrar_entrada_historial(entrada["id"])
                        st.rerun()
        if st.button("Vaciar historial completo"):
            borrar_historial_completo()
            st.rerun()


# ====================================
# FUNCIONES AUXILIARES DE LECTURA / CALIDAD DE DATOS
# ====================================

def normalizar_nombre(col: str) -> str:
    """Normaliza un nombre de columna para comparaciones flexibles."""
    return str(col).strip().lower().replace("_", " ")


def limpiar_encabezados(df: pd.DataFrame) -> pd.DataFrame:
    """Quita espacios sobrantes en los nombres de columnas."""
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    return df


def detectar_fila_encabezado(archivo_bytes: bytes, es_csv: bool, hoja=None, max_filas_revisar: int = 10):
    """
    Busca la primera fila que parece un encabezado real cuando hay filas
    en blanco antes del contenido. Devuelve el índice de fila (0-based).
    """
    try:
        if es_csv:
            vista = pd.read_csv(BytesIO(archivo_bytes), header=None, nrows=max_filas_revisar)
        else:
            vista = pd.read_excel(BytesIO(archivo_bytes), header=None, sheet_name=hoja, nrows=max_filas_revisar)
    except Exception:
        return 0

    for i in range(len(vista)):
        fila = vista.iloc[i]
        no_nulos = fila.notna().sum()
        if no_nulos >= max(2, int(0.6 * len(fila))):
            return i
    return 0


def intentar_convertir_numericas(df: pd.DataFrame) -> tuple[pd.DataFrame, list]:
    """Convierte a numéricas las columnas de texto que en realidad contienen números."""
    df = df.copy()
    convertidas = []
    for col in df.select_dtypes(include="object").columns:
        serie_convertida = pd.to_numeric(
            df[col].astype(str).str.replace(",", "", regex=False).str.strip(),
            errors="coerce"
        )
        proporcion_valida = serie_convertida.notna().mean()
        if proporcion_valida >= 0.8 and serie_convertida.notna().sum() > 0:
            df[col] = serie_convertida
            convertidas.append(col)
    return df, convertidas


def intentar_convertir_fechas(df: pd.DataFrame) -> tuple[pd.DataFrame, list]:
    """Detecta columnas que parecen fechas (por nombre o contenido) y las convierte."""
    df = df.copy()
    convertidas = []
    for col in df.columns:
        nombre_norm = normalizar_nombre(col)
        parece_fecha_por_nombre = any(p in nombre_norm for p in ["fecha", "date"])
        if df[col].dtype == "object" or parece_fecha_por_nombre:
            serie_convertida = pd.to_datetime(df[col], errors="coerce", dayfirst=True)
            proporcion_valida = serie_convertida.notna().mean()
            if proporcion_valida >= 0.7 and serie_convertida.notna().sum() > 0:
                df[col] = serie_convertida
                convertidas.append(col)
    return df, convertidas


def validar_calidad_datos(df: pd.DataFrame) -> list:
    """Revisa el dataframe y devuelve una lista de advertencias descriptivas."""
    advertencias = []

    if df.empty:
        advertencias.append("El archivo no contiene registros (está vacío).")
        return advertencias

    filas_totalmente_vacias = df.isna().all(axis=1).sum()
    if filas_totalmente_vacias > 0:
        advertencias.append(f"Se encontraron {filas_totalmente_vacias} fila(s) completamente vacía(s).")

    for col in df.columns:
        n_nulos = df[col].isna().sum()
        if n_nulos > 0:
            advertencias.append(f"La columna '{col}' tiene {n_nulos} valor(es) vacío(s).")

        if pd.api.types.is_numeric_dtype(df[col]):
            n_negativos = (df[col] < 0).sum()
            if n_negativos > 0:
                advertencias.append(f"La columna '{col}' tiene {n_negativos} valor(es) negativo(s).")

        if pd.api.types.is_datetime64_any_dtype(df[col]):
            hoy = pd.Timestamp.now()
            n_futuras = (df[col] > hoy).sum()
            if n_futuras > 0:
                advertencias.append(f"La columna '{col}' tiene {n_futuras} fecha(s) posterior(es) a hoy (posible inconsistencia).")

    return advertencias


# ====================================
# LECTORES MULTI-FORMATO (Excel, CSV, Word, PDF, SQLite)
# ====================================

def obtener_sub_opciones(nombre: str, info: dict):
    """
    Devuelve la lista de 'sub-elementos' seleccionables dentro de un archivo
    (hojas de Excel, tablas de Word/PDF, tablas de una base de datos SQLite).
    Devuelve None si el archivo es de un solo dataset (CSV) o si no se pudo inspeccionar.
    """
    tipo = info["tipo"]
    b = info["bytes"]
    try:
        if tipo in ("xlsx", "xls"):
            xls = pd.ExcelFile(BytesIO(b))
            return xls.sheet_names

        if tipo == "csv":
            return None

        if tipo == "docx":
            if not DOCX_DISPONIBLE:
                return None
            doc = docx.Document(BytesIO(b))
            etiquetas = [f"Tabla {i + 1}" for i in range(len(doc.tables))]
            etiquetas.append("Texto completo (párrafos)")
            return etiquetas

        if tipo == "pdf":
            if not PDFPLUMBER_DISPONIBLE:
                return None
            etiquetas = []
            with pdfplumber.open(BytesIO(b)) as pdf:
                for i, pagina in enumerate(pdf.pages):
                    n_tablas = len(pagina.extract_tables())
                    for j in range(n_tablas):
                        etiquetas.append(f"Página {i + 1} - Tabla {j + 1}")
            etiquetas.append("Texto completo (por página)")
            return etiquetas

        if tipo in ("db", "sqlite", "sqlite3"):
            ruta_tmp = Path(f"._tmp_{nombre}.sqlite")
            ruta_tmp.write_bytes(b)
            try:
                conn = sqlite3.connect(ruta_tmp)
                tablas = pd.read_sql(
                    "SELECT name FROM sqlite_master WHERE type='table'", conn
                )["name"].tolist()
                conn.close()
            finally:
                ruta_tmp.unlink(missing_ok=True)
            return tablas
    except Exception:
        return None
    return None


def leer_sub_dataset(nombre: str, info: dict, etiqueta, fila_encabezado: int = 0) -> pd.DataFrame:
    """Lee un dataset concreto (hoja/tabla/archivo completo) según el tipo de archivo."""
    tipo = info["tipo"]
    b = info["bytes"]

    if tipo == "csv":
        return pd.read_csv(BytesIO(b), header=fila_encabezado)

    if tipo in ("xlsx", "xls"):
        return pd.read_excel(BytesIO(b), sheet_name=etiqueta, header=fila_encabezado)

    if tipo == "docx":
        doc = docx.Document(BytesIO(b))
        if etiqueta == "Texto completo (párrafos)":
            parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
            return pd.DataFrame({"Texto": parrafos})
        idx = int(etiqueta.replace("Tabla ", "")) - 1
        tabla = doc.tables[idx]
        filas = [[c.text.strip() for c in fila.cells] for fila in tabla.rows]
        if len(filas) > 1:
            return pd.DataFrame(filas[1:], columns=filas[0])
        return pd.DataFrame(filas)

    if tipo == "pdf":
        if etiqueta == "Texto completo (por página)":
            textos = []
            with pdfplumber.open(BytesIO(b)) as pdf:
                for i, pagina in enumerate(pdf.pages):
                    textos.append({"Pagina": i + 1, "Texto": pagina.extract_text() or ""})
            return pd.DataFrame(textos)
        partes = etiqueta.replace("Página ", "").split(" - Tabla ")
        num_pagina = int(partes[0]) - 1
        num_tabla = int(partes[1]) - 1
        with pdfplumber.open(BytesIO(b)) as pdf:
            tablas = pdf.pages[num_pagina].extract_tables()
            tabla = tablas[num_tabla]
        if len(tabla) > 1:
            return pd.DataFrame(tabla[1:], columns=tabla[0])
        return pd.DataFrame(tabla)

    if tipo in ("db", "sqlite", "sqlite3"):
        ruta_tmp = Path(f"._tmp_{nombre}.sqlite")
        ruta_tmp.write_bytes(b)
        try:
            conn = sqlite3.connect(ruta_tmp)
            df = pd.read_sql(f'SELECT * FROM "{etiqueta}"', conn)
            conn.close()
        finally:
            ruta_tmp.unlink(missing_ok=True)
        return df

    raise ValueError(f"Tipo de archivo no soportado: {tipo}")


# ====================================
# EXPORTACIÓN (Excel / PDF con gráficos incluidos)
# ====================================

PALETA_COLORES = px.colors.qualitative.Plotly  # paleta de colores explícita para forzar color en la exportación

# Tamaño de embebido de cada gráfico en la hoja "Graficos" del Excel (en píxeles).
# Se fija explícitamente para que las imágenes no se dibujen "a tamaño completo" (lo que
# provocaba que unas taparan a otras); a partir de este tamaño se calcula cuántas filas
# hay que reservar para que la siguiente imagen no se superponga con la anterior.
ANCHO_IMG_EXCEL = 640
ALTO_IMG_EXCEL = 356
ALTO_FILA_EXCEL_PX = 20  # alto aproximado, en píxeles, de una fila por defecto en Excel
FILAS_POR_IMAGEN_EXCEL = (ALTO_IMG_EXCEL // ALTO_FILA_EXCEL_PX) + 4  # + margen para el título y separación


def fig_a_imagen_png(fig, ancho=900, alto=500):
    """
    Convierte una figura de Plotly a bytes PNG (requiere kaleido).
    Se fuerza explícitamente una plantilla con color (plotly_white) porque, si la
    figura queda sin plantilla resuelta, kaleido puede exportarla solo en negro.
    """
    if not KALEIDO_DISPONIBLE:
        return None
    try:
        fig = fig.update_layout(
            template="plotly_white",
            paper_bgcolor="white",
            plot_bgcolor="white",
            font=dict(color="black"),
        )
        return fig.to_image(format="png", width=ancho, height=alto, engine="kaleido", scale=2)
    except Exception:
        return None


def _dibujar_tabla_pdf(pdf: "FPDF", df: pd.DataFrame, max_columnas: int = 8, alto_fila: int = 6):
    """Dibuja una tabla completa en el PDF, paginando automáticamente (sin límite artificial de filas)."""
    columnas = list(df.columns)[:max_columnas]
    ancho_col = 190 / max(len(columnas), 1)

    def encabezado():
        pdf.set_font("Helvetica", "B", 8)
        for col in columnas:
            pdf.cell(ancho_col, alto_fila, str(col)[:15], border=1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)

    encabezado()
    for _, fila in df[columnas].iterrows():
        # Si la fila no cabe en la página actual, se agrega una nueva página y se repite el encabezado
        if pdf.get_y() + alto_fila > pdf.h - pdf.b_margin:
            pdf.add_page()
            encabezado()
        for col in columnas:
            pdf.cell(ancho_col, alto_fila, str(fila[col])[:15], border=1)
        pdf.ln()

    if len(df.columns) > max_columnas:
        pdf.set_font("Helvetica", "I", 7)
        pdf.cell(0, 6, f"(Se muestran {max_columnas} de {len(df.columns)} columnas)", ln=True)
        pdf.ln(2)


def _dibujar_grafico_pdf(pdf: "FPDF", titulo: str, img_bytes: bytes):
    """Agrega una página nueva con un gráfico (título + imagen)."""
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, titulo, ln=True)
    pdf.ln(2)
    try:
        # No se fija "y" a propósito: dejando que FPDF use la posición actual del
        # cursor, después de dibujar la imagen también actualiza pdf.get_y() a la
        # altura justo debajo de ella. Si se fija "y" manualmente (como antes), el
        # cursor interno no avanza y el contenido siguiente (el próximo dataset)
        # se dibuja encima de la imagen en vez de debajo.
        pdf.image(BytesIO(img_bytes), x=10, w=190)
        pdf.ln(6)
    except Exception:
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, "(No fue posible incrustar este gráfico)", ln=True)


def generar_pdf(datasets_reporte: dict, imagenes_por_dataset: dict, imagenes_sin_dataset: list | None = None) -> bytes:
    """
    Genera un PDF con, para CADA dataset incluido en el reporte:
    sus KPIs, la TABLA COMPLETA de datos filtrados (paginada, sin límite de filas) y,
    justo después de la tabla, el/los gráfico(s) que usan ese dataset (cada uno en su
    propia página, para que se vea grande y legible).

    datasets_reporte: {nombre_dataset: {"df": DataFrame_filtrado, "kpis": {...}}}
    imagenes_por_dataset: {nombre_dataset: [(titulo, bytes_png), ...]}
    imagenes_sin_dataset: gráficos "huérfanos" (caso excepcional, por si su dataset ya no
        está entre los seleccionados); se agregan al final para no perderlos.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Reporte de Datos", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 6, f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
    pdf.ln(2)

    for nombre_dataset, contenido in datasets_reporte.items():
        df = contenido["df"]
        kpis = contenido.get("kpis", {})

        if pdf.get_y() > 250:
            pdf.add_page()

        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, f"Dataset: {nombre_dataset}", ln=True)

        if kpis:
            pdf.set_font("Helvetica", "", 10)
            for etiqueta, valor in kpis.items():
                pdf.cell(0, 7, f"{etiqueta}: {valor}", ln=True)
            pdf.ln(2)

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, f"Registros filtrados: {len(df)}", ln=True)

        if df.empty:
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(0, 6, "(Sin registros con los filtros aplicados)", ln=True)
        else:
            _dibujar_tabla_pdf(pdf, df)

        pdf.ln(4)

        # Gráficos de ESTE dataset, justo después de su tabla (uno por página)
        for titulo, img_bytes in imagenes_por_dataset.get(nombre_dataset, []):
            _dibujar_grafico_pdf(pdf, titulo, img_bytes)

    # Por si quedara algún gráfico sin dataset asociado en el reporte (caso excepcional)
    for titulo, img_bytes in (imagenes_sin_dataset or []):
        _dibujar_grafico_pdf(pdf, titulo, img_bytes)

    return bytes(pdf.output(dest="S"))


def _nombre_hoja_valido(nombre: str, usados: set) -> str:
    """Excel limita los nombres de hoja a 31 caracteres y no permite ciertos símbolos; evita duplicados."""
    limpio = "".join(c for c in nombre if c not in r'[]:*?/\\')[:31] or "Datos"
    candidato = limpio
    contador = 1
    while candidato in usados:
        sufijo = f"_{contador}"
        candidato = f"{limpio[: 31 - len(sufijo)]}{sufijo}"
        contador += 1
    usados.add(candidato)
    return candidato


def generar_excel(datasets_reporte: dict, imagenes_graficos: list) -> bytes:
    """
    Genera un Excel con UNA HOJA POR CADA dataset incluido en el reporte (datos filtrados),
    una hoja "Indicadores" combinada y una hoja "Graficos" con los gráficos incrustados,
    cada uno con un tamaño fijo y suficiente espacio vertical para que no se superpongan.

    datasets_reporte: {nombre_dataset: {"df": DataFrame_filtrado, "kpis": {...}}}
    imagenes_graficos: [(titulo, bytes_png, nombre_dataset), ...]
    """
    buffer = BytesIO()
    nombres_hoja_usados = set()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        filas_kpis = []
        for nombre_dataset, contenido in datasets_reporte.items():
            hoja = _nombre_hoja_valido(nombre_dataset, nombres_hoja_usados)
            contenido["df"].to_excel(writer, sheet_name=hoja, index=False)
            for etiqueta, valor in contenido.get("kpis", {}).items():
                filas_kpis.append({"Dataset": nombre_dataset, "Indicador": etiqueta, "Valor": valor})

        if filas_kpis:
            pd.DataFrame(filas_kpis).to_excel(writer, sheet_name="Indicadores", index=False)

        if imagenes_graficos and OPENPYXL_IMAGENES_DISPONIBLE:
            workbook = writer.book
            hoja_graficos = workbook.create_sheet("Graficos")
            hoja_graficos.column_dimensions["A"].width = 90
            fila_actual = 1
            for titulo, img_bytes, _nombre_dataset in imagenes_graficos:
                hoja_graficos.cell(row=fila_actual, column=1, value=titulo)
                try:
                    img = OpenpyxlImage(BytesIO(img_bytes))
                    # Tamaño fijo para que todos los gráficos se vean parejos y no
                    # se superpongan entre sí (antes se insertaban a tamaño original).
                    img.width = ANCHO_IMG_EXCEL
                    img.height = ALTO_IMG_EXCEL
                    hoja_graficos.add_image(img, f"A{fila_actual + 1}")
                except Exception:
                    hoja_graficos.cell(row=fila_actual + 1, column=1, value="(No fue posible incrustar este gráfico)")
                fila_actual += FILAS_POR_IMAGEN_EXCEL

    return buffer.getvalue()


# ====================================
# CARGA DE ARCHIVOS (uno o varios, distintos formatos)
# ====================================

st.subheader("1. Cargar archivos")

archivos_subidos = st.file_uploader(
    "Seleccione uno o varios archivos (puede arrastrar varios a la vez, o agregar más después)",
    type=["xlsx", "xls", "csv", "docx", "pdf", "db", "sqlite", "sqlite3"],
    accept_multiple_files=True,
)

if archivos_subidos:
    for archivo in archivos_subidos:
        if archivo.name in st.session_state.archivos_originales:
            continue
        contenido = archivo.getvalue()
        if len(contenido) == 0:
            st.warning(f"El archivo '{archivo.name}' está vacío (0 bytes) y no se cargó.")
            continue
        extension = archivo.name.lower().split(".")[-1]
        st.session_state.archivos_originales[archivo.name] = {
            "bytes": contenido,
            "tipo": extension,
        }
    guardar_sesion()

if not st.session_state.archivos_originales:
    st.info("👆 Cargue uno o varios archivos (Excel, CSV, Word, PDF o SQLite) para comenzar.")
    st.stop()

# ====================================
# SELECCIÓN DE HOJAS / TABLAS POR ARCHIVO -> CONSTRUCCIÓN DE DATASETS
# ====================================

st.subheader("2. Archivos cargados")

nuevos_datasets = {}

for nombre, info in list(st.session_state.archivos_originales.items()):
    with st.expander(f"📄 {nombre}  ({info['tipo'].upper()})", expanded=True):
        col_info, col_quitar = st.columns([5, 1])
        with col_quitar:
            if st.button("🗑️ Quitar", key=f"quitar_{nombre}"):
                del st.session_state.archivos_originales[nombre]
                st.session_state.datasets_seleccionados = [
                    d for d in st.session_state.datasets_seleccionados if not d.startswith(nombre)
                ]
                guardar_sesion()
                st.rerun()

        tipo = info["tipo"]

        if tipo == "csv":
            fila_sugerida = detectar_fila_encabezado(info["bytes"], True)
            fila = st.number_input(
                "Fila donde comienza el encabezado real (0 = primera fila)",
                min_value=0, max_value=50, value=int(fila_sugerida), key=f"fila_{nombre}"
            )
            try:
                df = leer_sub_dataset(nombre, info, None, fila)
                df = limpiar_encabezados(df)
                df = df.loc[:, ~df.columns.astype(str).str.match(r"^Unnamed")]
                nuevos_datasets[nombre] = df
                st.caption(f"✅ {df.shape[0]} filas, {df.shape[1]} columnas")
            except Exception as e:
                st.error(f"No se pudo leer '{nombre}': {e}")

        else:
            sub_opciones = obtener_sub_opciones(nombre, info)

            if not sub_opciones:
                st.warning(
                    "No se pudieron detectar hojas/tablas en este archivo, "
                    "o falta instalar la librería necesaria para este formato."
                )
                continue

            if tipo in ("xlsx", "xls"):
                etiqueta_selector = "Hoja(s) a cargar"
            elif tipo in ("db", "sqlite", "sqlite3"):
                etiqueta_selector = "Tabla(s) a cargar"
            else:
                etiqueta_selector = "Elemento(s) a cargar (tabla o texto)"

            seleccionadas = st.multiselect(
                etiqueta_selector, sub_opciones, default=[sub_opciones[0]], key=f"sel_{nombre}"
            )

            for etiqueta in seleccionadas:
                fila = 0
                if tipo in ("xlsx", "xls"):
                    fila_sugerida = detectar_fila_encabezado(info["bytes"], False, etiqueta)
                    fila = st.number_input(
                        f"Fila de encabezado — {etiqueta}",
                        min_value=0, max_value=50, value=int(fila_sugerida),
                        key=f"fila_{nombre}_{etiqueta}"
                    )
                try:
                    df = leer_sub_dataset(nombre, info, etiqueta, fila)
                    df = limpiar_encabezados(df)
                    df = df.loc[:, ~df.columns.astype(str).str.match(r"^Unnamed")]
                    nombre_dataset = f"{nombre} [{etiqueta}]"
                    nuevos_datasets[nombre_dataset] = df
                    st.caption(f"✅ {etiqueta}: {df.shape[0]} filas, {df.shape[1]} columnas")
                except Exception as e:
                    st.error(f"No se pudo leer '{etiqueta}' de '{nombre}': {e}")

st.session_state.datasets = nuevos_datasets

if not st.session_state.datasets:
    st.info("No hay datasets disponibles todavía. Revise la selección de hojas/tablas arriba.")
    st.stop()

# Aplicar conversiones automáticas a cada dataset
for nombre_ds, df_ds in list(st.session_state.datasets.items()):
    df_ds, _ = intentar_convertir_numericas(df_ds)
    df_ds, _ = intentar_convertir_fechas(df_ds)
    st.session_state.datasets[nombre_ds] = df_ds

guardar_sesion()

# ====================================
# SELECCIÓN DE DATASETS PARA EL ANÁLISIS Y EL REPORTE (soporta varios a la vez)
# ====================================

st.subheader("3. Datasets a incluir en el análisis y el reporte")

nombres_datasets = list(st.session_state.datasets.keys())

seleccion_previa = [d for d in st.session_state.datasets_seleccionados if d in nombres_datasets]
if not seleccion_previa:
    seleccion_previa = [nombres_datasets[0]]

datasets_seleccionados = st.multiselect(
    "Elija uno o varios datasets (por ejemplo, provenientes de 2 archivos distintos). "
    "Cada uno se filtra por separado más abajo, y el reporte final combina los datos "
    "FILTRADOS de todos los que elija aquí.",
    nombres_datasets,
    default=seleccion_previa,
)

if not datasets_seleccionados:
    st.warning("Seleccione al menos un dataset para continuar.")
    st.stop()

st.session_state.datasets_seleccionados = datasets_seleccionados

# resultados_por_dataset: nombre_dataset -> {"df": DataFrame_filtrado, "kpis": {...}}
resultados_por_dataset = {}

tabs = st.tabs(datasets_seleccionados)

for nombre_ds, tab in zip(datasets_seleccionados, tabs):
    with tab:
        datos = st.session_state.datasets[nombre_ds]

        advertencias = validar_calidad_datos(datos)
        if advertencias:
            with st.expander(f"⚠️ Advertencias de calidad de datos ({len(advertencias)})", expanded=False):
                for a in advertencias:
                    st.warning(a)

        st.markdown("**Vista previa**")
        st.dataframe(datos, use_container_width=True)

        # --- Filtros propios de este dataset ---
        st.markdown("**Filtros**")
        datos_filtrados = datos.copy()

        columnas_categoricas = [
            c for c in datos.columns
            if (datos[c].dtype == "object" or str(datos[c].dtype) == "category")
            and datos[c].nunique(dropna=True) <= 50
        ]

        if columnas_categoricas:
            columnas_filtro_ui = st.columns(min(3, len(columnas_categoricas)))
            for i, col in enumerate(columnas_categoricas):
                valores = sorted(datos[col].dropna().unique().tolist())
                with columnas_filtro_ui[i % len(columnas_filtro_ui)]:
                    seleccionados = st.multiselect(
                        f"{col}", valores, default=valores, key=f"filtro_{nombre_ds}_{col}"
                    )
                if seleccionados:
                    datos_filtrados = datos_filtrados[datos_filtrados[col].isin(seleccionados)]
                else:
                    datos_filtrados = datos_filtrados.iloc[0:0]
        else:
            st.caption("No se detectaron columnas categóricas para filtrar.")

        columnas_fecha = [c for c in datos.columns if pd.api.types.is_datetime64_any_dtype(datos[c])]
        for col in columnas_fecha:
            col_valida = datos[col].dropna()
            if col_valida.empty:
                continue
            fecha_min, fecha_max = col_valida.min().date(), col_valida.max().date()
            if fecha_min == fecha_max:
                continue
            rango = st.date_input(
                f"Rango de {col}", value=(fecha_min, fecha_max), key=f"fecha_{nombre_ds}_{col}"
            )
            if isinstance(rango, tuple) and len(rango) == 2:
                datos_filtrados = datos_filtrados[
                    (datos_filtrados[col].dt.date >= rango[0]) & (datos_filtrados[col].dt.date <= rango[1])
                ]

        if datos_filtrados.empty:
            st.warning("Los filtros seleccionados no arrojan ningún registro para este dataset.")

        # --- Indicadores propios de este dataset ---
        st.markdown("**Indicadores**")
        columnas_numericas = datos_filtrados.select_dtypes(include="number").columns.tolist()
        kpis_dataset = {}

        if columnas_numericas and not datos_filtrados.empty:
            columna_indicador = st.selectbox(
                "Columna numérica para indicadores", columnas_numericas, key=f"kpi_col_{nombre_ds}"
            )
            total = datos_filtrados[columna_indicador].sum()
            promedio = datos_filtrados[columna_indicador].mean()
            registros = len(datos_filtrados)

            c1, c2, c3 = st.columns(3)
            c1.metric("Total", f"{total:,.2f}")
            c2.metric("Promedio", f"{promedio:,.2f}")
            c3.metric("Registros", registros)

            kpis_dataset = {
                "Columna analizada": columna_indicador,
                "Total": f"{total:,.2f}",
                "Promedio": f"{promedio:,.2f}",
                "Registros": registros,
            }
        else:
            st.caption("No hay columnas numéricas disponibles, o los filtros actuales no dejan registros.")

        st.markdown("**Datos filtrados**")
        st.dataframe(datos_filtrados, use_container_width=True)

        resultados_por_dataset[nombre_ds] = {"df": datos_filtrados, "kpis": kpis_dataset}

# ====================================
# GENERADOR DE GRÁFICOS (usa los datos YA FILTRADOS de cada dataset seleccionado)
# ====================================

st.subheader("4. Generador de Gráficos")

dataset_para_grafico = st.selectbox(
    "Dataset a usar en el nuevo gráfico (se grafican los datos ya filtrados arriba)",
    datasets_seleccionados,
    key="selector_dataset_nuevo_grafico"
)
df_para_grafico = resultados_por_dataset[dataset_para_grafico]["df"]
columnas_grafico = df_para_grafico.columns.tolist()
columnas_num_grafico = df_para_grafico.select_dtypes(include="number").columns.tolist()

if df_para_grafico.empty:
    st.info(f"El dataset '{dataset_para_grafico}' no tiene registros con los filtros actuales.")
elif not columnas_num_grafico:
    st.info(f"El dataset '{dataset_para_grafico}' no tiene columnas numéricas para graficar.")
else:
    with st.form("form_nuevo_grafico"):
        col_x, col_y, col_tipo = st.columns(3)
        with col_x:
            eje_x_nuevo = st.selectbox("Eje X", columnas_grafico, key="nuevo_eje_x")
        with col_y:
            eje_y_nuevo = st.selectbox("Eje Y (numérico)", columnas_num_grafico, key="nuevo_eje_y")
        with col_tipo:
            tipo_nuevo = st.selectbox(
                "Tipo de gráfico",
                ["Barras", "Pastel", "Líneas", "Dispersión", "Histograma", "Boxplot"],
                key="nuevo_tipo"
            )
        agregar = st.form_submit_button("➕ Agregar gráfico")
        if agregar:
            st.session_state.graficos.append({
                "id": f"g{len(st.session_state.graficos)}_{datetime.now().timestamp()}",
                "dataset": dataset_para_grafico,
                "eje_x": eje_x_nuevo,
                "eje_y": eje_y_nuevo,
                "tipo": tipo_nuevo,
            })
            guardar_sesion()
            st.rerun()

# Cada elemento: (titulo, bytes_png, nombre_dataset) — se conserva el dataset de origen
# para poder agrupar los gráficos por dataset al exportar (PDF) o para dar contexto (Excel).
imagenes_graficos = []

if not st.session_state.graficos:
    st.info("Aún no ha agregado gráficos. Use el formulario de arriba para crear el primero.")
else:
    for indice_color, g in enumerate(list(st.session_state.graficos)):
        contenido_ds = resultados_por_dataset.get(g["dataset"])
        if contenido_ds is None:
            st.warning(
                f"El dataset '{g['dataset']}' de este gráfico no está entre los seleccionados en el paso 3. "
                "Agréguelo arriba para poder verlo y exportarlo."
            )
            continue

        df_g = contenido_ds["df"]
        if df_g.empty:
            st.info(f"'{g['dataset']}' no tiene registros con los filtros actuales; este gráfico se omite.")
            continue

        if len(df_g) > 50000:
            st.warning(
                f"⚠️ El dataset tiene {len(df_g):,} registros. Es muy grande y el gráfico puede tardar en cargar."
            )

        st.markdown(f"**{g['tipo']}** — {g['eje_y']} por {g['eje_x']}  _(dataset: {g['dataset']}, datos filtrados)_")

        try:
            if g["eje_x"] == g["eje_y"]:
                resumen = df_g[[g["eje_x"]]].groupby(g["eje_x"]).size().reset_index(name="Cantidad")
                col_valor = "Cantidad"
                st.caption("Mismo campo en ambos ejes: se muestra el conteo de registros.")
            else:
                resumen = df_g.groupby(g["eje_x"])[g["eje_y"]].sum().reset_index()
                col_valor = g["eje_y"]

            color_principal = PALETA_COLORES[indice_color % len(PALETA_COLORES)]

            if g["tipo"] == "Barras":
                fig = px.bar(
                    resumen,
                    x=g["eje_x"],
                    y=col_valor,
                    title=f"{col_valor} por {g['eje_x']}",
                    color=g["eje_x"],
                    color_discrete_sequence=PALETA_COLORES,
                )

            elif g["tipo"] == "Pastel":
                fig = px.pie(
                    resumen,
                    names=g["eje_x"],
                    values=col_valor,
                    title=f"Distribución de {col_valor}",
                    color_discrete_sequence=PALETA_COLORES,
                )

            elif g["tipo"] == "Líneas":
                fig = px.line(
                    resumen,
                    x=g["eje_x"],
                    y=col_valor,
                    title=f"{col_valor} por {g['eje_x']}",
                    markers=True,
                    color_discrete_sequence=[color_principal],
                )

            elif g["tipo"] == "Histograma":
                fig = px.histogram(
                    df_g,
                    x=g["eje_y"],
                    title=f"Histograma de {g['eje_y']}",
                    color_discrete_sequence=[color_principal],
                )

            elif g["tipo"] == "Boxplot":
                fig = px.box(
                    df_g,
                    x=g["eje_x"],
                    y=g["eje_y"],
                    title=f"Boxplot de {g['eje_y']} por {g['eje_x']}",
                    color_discrete_sequence=[color_principal],
                )

            else:
                fig = px.scatter(
                    df_g,
                    x=g["eje_x"],
                    y=g["eje_y"],
                    title=f"{g['eje_y']} vs {g['eje_x']}",
                    color_discrete_sequence=[color_principal],
                )
            st.plotly_chart(fig, use_container_width=True, key=f"chart_{g['id']}")

            imagen_png = fig_a_imagen_png(fig)
            if imagen_png:
                titulo_grafico = f"{g['tipo']} - {col_valor} por {g['eje_x']} ({g['dataset']})"
                imagenes_graficos.append((titulo_grafico, imagen_png, g["dataset"]))

        except Exception as e:
            st.error(f"No fue posible generar este gráfico: {e}")

        if st.button("🗑️ Eliminar este gráfico", key=f"del_{g['id']}"):
            st.session_state.graficos = [x for x in st.session_state.graficos if x["id"] != g["id"]]
            guardar_sesion()
            st.rerun()

        st.divider()

if st.session_state.graficos and not KALEIDO_DISPONIBLE:
    st.caption(
        "⚠️ Para incluir los gráficos en los reportes de Excel y PDF, instale: `pip install -U kaleido`"
    )

# Agrupar los gráficos por dataset, para que en el PDF cada uno aparezca justo
# debajo de la tabla de datos filtrados de su dataset correspondiente.
imagenes_por_dataset = defaultdict(list)
imagenes_sin_dataset = []
for titulo, img_bytes, nombre_dataset_grafico in imagenes_graficos:
    if nombre_dataset_grafico in resultados_por_dataset:
        imagenes_por_dataset[nombre_dataset_grafico].append((titulo, img_bytes))
    else:
        imagenes_sin_dataset.append((titulo, img_bytes))

# ====================================
# EXPORTAR REPORTE (Excel, CSV y PDF — combina TODOS los datasets seleccionados, ya filtrados)
# ====================================

st.subheader("5. Exportar Reporte")

st.caption(
    "El reporte combina los datos FILTRADOS de los datasets elegidos en el paso 3: "
    + ", ".join(datasets_seleccionados)
)

col_excel, col_csv, col_pdf = st.columns(3)

with col_excel:
    try:
        buffer_excel = generar_excel(resultados_por_dataset, imagenes_graficos)
        st.download_button(
            label="📥 Descargar Excel",
            data=buffer_excel,
            file_name="Reporte.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        st.caption(f"Una hoja por dataset ({len(resultados_por_dataset)}).")
        if imagenes_graficos:
            st.caption(f"Incluye {len(imagenes_graficos)} gráfico(s) en la hoja 'Graficos', uno debajo del otro sin superponerse.")
    except Exception as e:
        st.error(f"No fue posible generar el Excel: {e}")

with col_csv:
    st.caption("El CSV no admite varias tablas: descargue un archivo por dataset.")
    for nombre_ds in datasets_seleccionados:
        df_export = resultados_por_dataset[nombre_ds]["df"]
        try:
            buffer_csv = df_export.to_csv(index=False).encode("utf-8-sig")
            nombre_archivo = "Reporte_" + "".join(c for c in nombre_ds if c.isalnum())[:40] + ".csv"
            st.download_button(
                label=f"📥 CSV: {nombre_ds}",
                data=buffer_csv,
                file_name=nombre_archivo,
                mime="text/csv",
                key=f"csv_{nombre_ds}",
            )
        except Exception as e:
            st.error(f"No fue posible generar el CSV de '{nombre_ds}': {e}")

with col_pdf:
    if not FPDF_DISPONIBLE:
        st.caption("Para exportar a PDF instale la librería: pip install fpdf2")
    else:
        try:
            pdf_bytes = generar_pdf(resultados_por_dataset, imagenes_por_dataset, imagenes_sin_dataset)
            st.download_button(
                label="📥 Descargar PDF",
                data=pdf_bytes,
                file_name="Reporte.pdf",
                mime="application/pdf"
            )
            total_registros = sum(len(c["df"]) for c in resultados_por_dataset.values())
            st.caption(f"Incluye {total_registros} registro(s) en total (todas las filas, paginado automático).")
            if imagenes_graficos:
                st.caption(f"Incluye {len(imagenes_graficos)} gráfico(s), cada uno justo después de la tabla de su dataset.")
        except Exception as e:
            st.error(f"No fue posible generar el PDF: {e}")

# Guardar la sesión al final de cada ejecución (archivos, gráficos, selección de datasets)
guardar_sesion()